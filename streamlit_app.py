import streamlit as st
import os
import warnings
import time
from datetime import datetime
import io
from io import BytesIO
warnings.filterwarnings('ignore')

# FIX FOR CHROMADB/SQLITE3
try:
    import sys
    import pysqlite3
    sys.modules['sqlite3'] = sys.modules.pop('pysqlite3')
    import sqlite3
    st.info(f"✅ SQLite patched to version: {sqlite3.sqlite_version}")
except ImportError:
    st.error("❌ pysqlite3-binary not installed. Add to requirements.txt.")

# Import Required Libraries
from crewai import Agent, Task, Crew, Process
from crewai.llm import LLM
import json

# DOCX Libraries for Word Document Generation
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn

# Fix RGB Color import
try:
    from docx.shared import RGBColor
except ImportError:
    pass  # Use default colors

# Configure Streamlit Page
st.set_page_config(
    page_title="Multi-Agent AI Analysis System",
    page_icon="🤖",
    layout="wide"
)

# Custom CSS for professional look
st.markdown("""
<style>
    .main-header {
        font-size: 2.5em;
        font-weight: bold;
        text-align: center;
        color: #1f77b4;
        margin-bottom: 0.5em;
    }
    .sub-header {
        font-size: 1.5em;
        font-weight: bold;
        color: #ff7f0e;
        margin-bottom: 1em;
    }
    .status-box {
        background-color: #f0f2f6;
        padding: 1em;
        border-radius: 0.5em;
        border-left: 5px solid #1f77b4;
    }
</style>
""", unsafe_allow_html=True)

# STEP 4: Configure Groq LLM with Rate Limiting
@st.cache_resource
def get_groq_llm(api_key):
    return LLM(
        model="groq/llama-3.1-8b-instant",
        api_key=api_key,
        max_tokens=1000,
        temperature=0.1
    )

# STEP 5: Professional Document Formatter Class (adapted for Streamlit)
class MultiAgentDocumentFormatter:
    """
    Professional Word Document Formatter for Multi-Agent Analysis Results
    Creates executive-level reports with proper formatting and structure
    """

    def __init__(self):
        self.document = Document()
        self.setup_styles()

    def setup_styles(self):
        """Setup custom styles for professional document formatting"""
        try:
            # Title Style
            title_style = self.document.styles.add_style('Custom Title', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Calibri'
            title_font.size = Pt(24)
            title_font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Heading Style
            heading_style = self.document.styles.add_style('Custom Heading', WD_STYLE_TYPE.PARAGRAPH)
            heading_font = heading_style.font
            heading_font.name = 'Calibri'
            heading_font.size = Pt(16)
            heading_font.bold = True

            # Body Style
            body_style = self.document.styles.add_style('Custom Body', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Calibri'
            body_font.size = Pt(11)

        except Exception as e:
            st.warning(f"Style setup note: {e}")

    def add_header(self, title, query, timestamp):
        """Add professional document header"""
        # Main Title
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(20)
        title_run.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add spacing
        self.document.add_paragraph()

        # Query Information
        query_label = self.document.add_paragraph()
        query_label.add_run("Analysis Query:").bold = True
        query_para = self.document.add_paragraph(f'"{query}"')
        query_para.italic = True

        # Timestamp
        time_label = self.document.add_paragraph()
        time_label.add_run("Generated:").bold = True
        self.document.add_paragraph(timestamp)

        # Analysis Information
        method_label = self.document.add_paragraph()
        method_label.add_run("Analysis Method:").bold = True
        self.document.add_paragraph("Hybrid Multi-Agent System (Research → Analysis → Synthesis)")

        # Add separator line
        self.document.add_paragraph("_" * 80)
        self.document.add_paragraph()

    def add_section(self, heading, content, is_main_section=True):
        """Add formatted section with heading and content"""
        if is_main_section:
            heading_para = self.document.add_paragraph()
            heading_run = heading_para.add_run(heading)
            heading_run.font.name = 'Calibri'
            heading_run.font.size = Pt(14)
            heading_run.bold = True
        else:
            heading_para = self.document.add_paragraph()
            heading_run = heading_para.add_run(heading)
            heading_run.font.name = 'Calibri'
            heading_run.font.size = Pt(12)
            heading_run.bold = True

        # Format content with proper paragraphs
        content_lines = str(content).split('\n')
        for line in content_lines:
            if line.strip():
                para = self.document.add_paragraph(line.strip())
                para.style.font.name = 'Calibri'
                para.style.font.size = Pt(11)

        # Add spacing after section
        self.document.add_paragraph()

    def add_executive_summary_box(self, summary):
        """Add highlighted executive summary box"""
        # Add Executive Summary Header
        summary_heading = self.document.add_paragraph()
        summary_run = summary_heading.add_run("🎯 EXECUTIVE SUMMARY")
        summary_run.font.name = 'Calibri'
        summary_run.font.size = Pt(14)
        summary_run.bold = True

        # Add summary content
        for line in str(summary).split('\n'):
            if line.strip():
                summary_para = self.document.add_paragraph(line.strip())
                summary_para.style.font.name = 'Calibri'
                summary_para.style.font.size = Pt(12)

        self.document.add_paragraph("_" * 80)
        self.document.add_paragraph()

    def get_document_bytes(self):
        """Return document as bytes for download"""
        bio = BytesIO()
        self.document.save(bio)
        bio.seek(0)
        return bio.getvalue()

# STEP 6: Define Specialized Agents with Rate Limiting
@st.cache_resource
def create_agents(api_key):
    groq_llm = get_groq_llm(api_key)

    # RESEARCH AGENT
    research_agent = Agent(
        role='Senior Research Specialist',
        goal='Conduct comprehensive research using available knowledge and provide structured findings',
        backstory="""You are an expert researcher with vast knowledge across multiple domains.
        You excel at analyzing questions, breaking them down into key components, and providing
        comprehensive information based on your extensive knowledge base. You organize findings
        clearly with proper sections and provide valuable context and insights.""",
        verbose=False,
        allow_delegation=False,
        llm=groq_llm,
        max_iter=2,
        max_rpm=10
    )

    # ANALYSIS AGENT
    analysis_agent = Agent(
        role='Senior Data Analyst and Strategic Advisor',
        goal='Perform comprehensive analysis and provide actionable insights and recommendations',
        backstory="""You are a world-class analyst with expertise in:
        - Statistical analysis and trend identification
        - Pattern recognition and anomaly detection
        - Financial analysis and risk assessment
        - Strategic thinking and decision support
        - Predictive analysis and forecasting
        - Text analytics and sentiment analysis
        You synthesize information into clear, actionable insights and strategic recommendations.""",
        verbose=False,
        allow_delegation=False,
        llm=groq_llm,
        max_iter=2,
        max_rpm=10
    )

    # ORCHESTRATOR AGENT
    orchestrator_agent = Agent(
        role='Chief Orchestrator and Synthesis Expert',
        goal='Coordinate workflows and synthesize comprehensive executive-level reports',
        backstory="""You are a master coordinator and executive advisor who excels at:
        - Managing complex multi-step analyses
        - Synthesizing diverse information sources
        - Creating executive-level summaries and reports
        - Providing strategic recommendations and roadmaps
        - Ensuring comprehensive coverage of all aspects
        - Risk assessment and implementation planning
        You deliver clear, well-structured, actionable final reports suitable for C-level executives.""",
        verbose=False,
        allow_delegation=False,
        llm=groq_llm,
        max_iter=2,
        max_rpm=10
    )

    return research_agent, analysis_agent, orchestrator_agent

# STEP 7: Define Enhanced Tasks for Document Formatting
def create_research_task(user_query):
    return Task(
        description=f"""
        Conduct comprehensive research and analysis on: "{user_query}"

        Your research must include the following structured sections:

        1. KEY FINDINGS:
        - Most important facts and current information
        - Critical statistics and quantitative data
        - Recent developments and updates

        2. CONTEXTUAL BACKGROUND:
        - Relevant historical context
        - Industry or domain background
        - Important stakeholders and players

        3. MULTIPLE PERSPECTIVES:
        - Different viewpoints and opinions
        - Contrasting approaches or solutions
        - Expert opinions and insights

        4. CURRENT TRENDS:
        - Latest developments and movements
        - Emerging patterns and shifts
        - Future indicators and signals

        Structure your response with clear section headers for professional document formatting.
        Keep each section concise but comprehensive to manage token limits effectively.
        """,
        agent=None,
        expected_output="Well-structured research report with 4 distinct sections: Key Findings, Contextual Background, Multiple Perspectives, and Current Trends (max 800 words)"
    )

def create_analysis_task(user_query, research_results=None):
    research_context = f"\n\nRESEARCH CONTEXT:\n{research_results}" if research_results else ""

    return Task(
        description=f"""
        Perform comprehensive multi-dimensional analysis on: "{user_query}"{research_context}

        Provide thorough analysis organized in these exact sections:

        1. QUANTITATIVE INSIGHTS:
        - Statistical analysis of numerical data
        - Trend identification and pattern recognition
        - Comparative analysis and benchmarking
        - Growth rates, percentages, and key metrics

        2. QUALITATIVE ANALYSIS:
        - Strategic implications and significance
        - Opportunities and potential benefits
        - Challenges and obstacles identified
        - Quality factors and subjective assessments

        3. RISK ASSESSMENT:
        - Potential risks and uncertainties
        - Probability and impact evaluation
        - Mitigation strategies and safeguards
        - Contingency considerations

        4. PREDICTIVE ANALYSIS:
        - Future trends and forecasting
        - Scenario planning and projections
        - Expected outcomes and timelines
        - Leading indicators to monitor

        5. STRATEGIC RECOMMENDATIONS:
        - Top 3-5 actionable recommendations
        - Priority ranking and implementation sequence
        - Resource requirements and success metrics
        - Expected ROI and impact assessment

        Use clear section headers and bullet points for professional document formatting.
        """,
        agent=None,
        expected_output="Comprehensive analysis with 5 sections: Quantitative Insights, Qualitative Analysis, Risk Assessment, Predictive Analysis, and Strategic Recommendations (max 800 words)"
    )

def create_orchestration_task(user_query, research_data, analysis_data):
    return Task(
        description=f"""
        Create an executive-level synthesis report for: "{user_query}"

        Synthesize and coordinate the following information:

        RESEARCH DATA:
        {research_data}

        ANALYSIS DATA:
        {analysis_data}

        Create a comprehensive executive report with these EXACT sections:

        1. EXECUTIVE SUMMARY:
        - 3-4 key sentences summarizing the entire analysis
        - Most critical insights and conclusions
        - Primary recommendation or course of action

        2. KEY FINDINGS:
        - Top 5 most important discoveries from research and analysis
        - Critical facts that drive decision-making
        - Validated insights with supporting evidence

        3. STRATEGIC RECOMMENDATIONS:
        - Top 3 priority actions ranked by importance
        - Clear implementation steps for each recommendation
        - Expected outcomes and success measures

        4. IMPLEMENTATION ROADMAP:
        - Specific next steps with timelines
        - Resource requirements and responsibilities
        - Milestones and checkpoints
        - Dependencies and prerequisites

        5. RISK CONSIDERATIONS:
        - Major risks and potential obstacles
        - Impact assessment and probability
        - Mitigation strategies and contingency plans

        6. CONCLUSION:
        - Final strategic assessment and outlook
        - Success factors and critical requirements
        - Long-term implications and considerations

        Ensure professional formatting with clear headers, bullet points, and executive-level language.
        """,
        agent=None,
        expected_output="Executive-level comprehensive report with 6 distinct sections: Executive Summary, Key Findings, Strategic Recommendations, Implementation Roadmap, Risk Considerations, and Conclusion"
    )

# STEP 8: Enhanced Multi-Agent System with DOCX Export
def run_hybrid_multi_agent_analysis_with_docx(user_query, api_key, export_to_docx=True):
    """
    Execute complete hybrid multi-agent workflow with professional DOCX export
    Returns results dict and optional docx bytes
    """

    status_placeholder = st.empty()
    status_placeholder.info(f"🚀 Starting Enhanced Hybrid Multi-Agent Analysis...")
    status_placeholder.info(f"📝 Query: {user_query}")
    timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    status_placeholder.info(f"⏰ Started at: {timestamp}")
    status_placeholder.info(f"📄 DOCX Export: {'Enabled' if export_to_docx else 'Disabled'}")
    status_placeholder.info(f"⚠️ Rate Limiting: Conservative 8 RPM per phase")

    try:
        research_agent, analysis_agent, orchestrator_agent = create_agents(api_key)

        # Phase 1: Research Agent
        status_placeholder.info("\n🔍 PHASE 1: Research Agent - Information Gathering")
        status_placeholder.info("Status: Conducting comprehensive research...")

        research_task = create_research_task(user_query)
        research_task.agent = research_agent
        research_crew = Crew(
            agents=[research_agent],
            tasks=[research_task],
            process=Process.sequential,
            verbose=False,
            max_rpm=8
        )

        with st.spinner('Research in progress...'):
            research_results = research_crew.kickoff()
        status_placeholder.success(f"✅ Research Phase Complete - {len(str(research_results))} characters generated")

        # Rate limiting delay
        status_placeholder.info("⏳ Rate limit management: Waiting 10 seconds...")
        time.sleep(10)

        # Phase 2: Analysis Agent
        status_placeholder.info("\n📊 PHASE 2: Analysis Agent - Comprehensive Analysis")
        status_placeholder.info("Status: Performing multi-dimensional analysis...")

        analysis_task = create_analysis_task(user_query, str(research_results))
        analysis_task.agent = analysis_agent
        analysis_crew = Crew(
            agents=[analysis_agent],
            tasks=[analysis_task],
            process=Process.sequential,
            verbose=False,
            max_rpm=8
        )

        with st.spinner('Analysis in progress...'):
            analysis_results = analysis_crew.kickoff()
        status_placeholder.success(f"✅ Analysis Phase Complete - {len(str(analysis_results))} characters generated")

        # Rate limiting delay
        status_placeholder.info("⏳ Rate limit management: Waiting 10 seconds...")
        time.sleep(10)

        # Phase 3: Orchestrator Agent
        status_placeholder.info("\n🎯 PHASE 3: Orchestrator Agent - Executive Synthesis")
        status_placeholder.info("Status: Creating executive-level comprehensive report...")

        orchestration_task = create_orchestration_task(user_query, str(research_results), str(analysis_results))
        orchestration_task.agent = orchestrator_agent
        orchestration_crew = Crew(
            agents=[orchestrator_agent],
            tasks=[orchestration_task],
            process=Process.sequential,
            verbose=False,
            max_rpm=8
        )

        with st.spinner('Synthesis in progress...'):
            final_results = orchestration_crew.kickoff()
        status_placeholder.success(f"✅ Orchestration Phase Complete - {len(str(final_results))} characters generated")

        status_placeholder.success("✅ HYBRID MULTI-AGENT ANALYSIS COMPLETE")
        status_placeholder.info(f"⏰ Completed at: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

        # Compile results
        results = {
            'query': user_query,
            'timestamp': timestamp,
            'research': research_results,
            'analysis': analysis_results,
            'final_report': final_results
        }

        # DOCX Export Phase
        docx_bytes = None
        if export_to_docx:
            status_placeholder.info("\n📄 PHASE 4: Professional Document Generation")
            status_placeholder.info("Status: Creating Word document with executive formatting...")
            try:
                docx_bytes = create_docx_bytes(results)
                results['docx_bytes'] = docx_bytes
                status_placeholder.success("✅ DOCX Export Complete")
            except Exception as e:
                status_placeholder.error(f"⚠️ DOCX Export Error: {e}")
                status_placeholder.info("📋 Analysis results still available as text output")

        return results, docx_bytes

    except Exception as e:
        error_msg = str(e).lower()
        if "rate_limit" in error_msg or "429" in error_msg:
            st.error(f"🚨 RATE LIMIT ERROR: {e}")
            st.info("""
            💡 SOLUTIONS:
            1. Wait 60 seconds and try again
            2. Use a shorter, simpler query
            3. Consider upgrading your Groq plan for higher limits
            """)
            return {"error": "rate_limit", "message": str(e)}, None
        else:
            st.error(f"🚨 SYSTEM ERROR: {e}")
            st.info("""
            💡 TROUBLESHOOTING:
            1. Check your internet connection
            2. Verify your Groq API key is valid
            3. Try restarting the app if issues persist
            """)
            return {"error": "system", "message": str(e)}, None

# STEP 9: Professional DOCX Report Generator
def create_docx_bytes(results):
    """
    Generate a professional Word document from multi-agent analysis results
    Returns bytes for download
    """

    try:
        st.info("Creating document structure...")
        formatter = MultiAgentDocumentFormatter()

        # Document Header with Title and Metadata
        formatter.add_header(
            title="MULTI-AGENT INTELLIGENCE ANALYSIS REPORT",
            query=results['query'],
            timestamp=results['timestamp']
        )

        # Extract and format Executive Summary
        final_report = str(results['final_report'])
        if "EXECUTIVE SUMMARY" in final_report.upper():
            try:
                summary_start = final_report.upper().find("EXECUTIVE SUMMARY")
                summary_content = final_report[summary_start:summary_start+500]
                lines = summary_content.split('\n')
                clean_summary = []
                for line in lines[1:6]:
                    if line.strip() and not line.upper().startswith(('KEY FINDINGS', 'STRATEGIC')):
                        clean_summary.append(line.strip())

                if clean_summary:
                    formatter.add_executive_summary_box(' '.join(clean_summary))
            except:
                formatter.add_executive_summary_box("Executive summary extracted from comprehensive analysis below.")

        # Research Findings Section
        formatter.add_section("🔍 RESEARCH FINDINGS & INTELLIGENCE", results['research'])

        # Analysis Section
        formatter.add_section("📊 COMPREHENSIVE ANALYSIS & INSIGHTS", results['analysis'])

        # Final Executive Report
        formatter.add_section("🎯 EXECUTIVE SYNTHESIS & RECOMMENDATIONS", results['final_report'])

        # System Information Footer
        system_info = f"""Analysis Framework: Hybrid Multi-Agent Intelligence System
Research Agent: Senior Research Specialist with domain expertise
Analysis Agent: Senior Data Analyst and Strategic Advisor
Orchestrator Agent: Chief Synthesis Expert and Executive Advisor
LLM Technology: Groq Llama-3.1-8B-Instant
Processing Method: Sequential multi-phase analysis with rate limiting
Document Generated: {results['timestamp']}
Total Processing Time: ~3-4 minutes including safety delays"""

        formatter.add_section("ℹ️ SYSTEM METADATA", system_info, is_main_section=False)

        # Generate timestamped filename
        safe_query = ''.join(c for c in results['query'][:30] if c.isalnum() or c in (' ', '-', '_')).strip()
        safe_query = safe_query.replace(' ', '_')
        timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"Multi_Agent_Analysis_{safe_query}_{timestamp_str}.docx"

        # Get bytes
        docx_bytes = formatter.get_document_bytes()
        st.success("✅ Professional Word document created successfully")
        return docx_bytes, filename

    except Exception as e:
        st.error(f"⚠️ DOCX Generation Error: {e}")
        st.info("📋 Analysis results are still available as text output below")
        st.info("💡 You can manually copy the text results if needed")
        return None, None

# Main Streamlit App
def main():
    st.markdown('<h1 class="main-header">🤖 Multi-Agent AI Analysis System</h1>', unsafe_allow_html=True)
    st.markdown('<p style="text-align: center; color: #666;">Professional Intelligence Analysis with Executive Reporting</p>', unsafe_allow_html=True)

    # Sidebar for API Key
    st.sidebar.header("🔑 Configuration")
    api_key = st.sidebar.text_input("Groq API Key", type="password", help="Enter your Groq API key (starts with 'gsk_')")
    # For production, uncomment the next line and use secrets
    # api_key = st.secrets.get("GROQ_API_KEY", "")

    export_docx = st.sidebar.checkbox("📄 Generate DOCX Report", value=True, help="Create professional Word document for download")

    # Demo Query Option
    use_demo = st.sidebar.checkbox("🎯 Use Demo Query", value=False)
    if use_demo:
        default_query = "What are the top 3 emerging AI trends for 2025 and their business impact?"
    else:
        default_query = ""

    # Main Query Input
    st.markdown('<h2 class="sub-header">📝 Enter Your Analysis Query</h2>', unsafe_allow_html=True)
    user_query = st.text_area(
        "What would you like to analyze?",
        value=default_query,
        height=100,
        placeholder="e.g., What are the key investment opportunities in renewable energy?"
    )

    # Run Button
    if st.button("🚀 Run Multi-Agent Analysis", type="primary", disabled=not api_key or not user_query.strip()):
        if not api_key:
            st.warning("⚠️ Please enter your Groq API key in the sidebar.")
            return
        if not user_query.strip():
            st.warning("⚠️ Please enter a query.")
            return

        # Progress Bar
        progress_bar = st.progress(0)
        status_text = st.empty()

        # Run Analysis
        with st.spinner("Initiating multi-agent workflow..."):
            results, docx_data = run_hybrid_multi_agent_analysis_with_docx(user_query, api_key, export_docx)

        if 'error' not in results:
            st.success("✅ Analysis Complete! View results below.")

            # Display Results
            col1, col2 = st.columns(2)

            with col1:
                st.markdown('<h3 class="sub-header">🔍 Research Findings</h3>', unsafe_allow_html=True)
                st.markdown(f'<div class="status-box">{results["research"]}</div>', unsafe_allow_html=True)

                st.markdown('<h3 class="sub-header">📊 Analysis Insights</h3>', unsafe_allow_html=True)
                st.markdown(f'<div class="status-box">{results["analysis"]}</div>', unsafe_allow_html=True)

            with col2:
                st.markdown('<h3 class="sub-header">🎯 Executive Report</h3>', unsafe_allow_html=True)
                st.markdown(f'<div class="status-box">{results["final_report"]}</div>', unsafe_allow_html=True)

            # Download DOCX
            if docx_data and export_docx:
                docx_bytes, filename = docx_data
                st.download_button(
                    label="📥 Download Professional DOCX Report",
                    data=docx_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download the executive-level Word document"
                )

            # Example Queries
            st.markdown("---")
            st.markdown('<h4 class="sub-header">💡 Example Queries to Try</h4>', unsafe_allow_html=True)
            examples = [
                "What are the key investment opportunities in renewable energy?",
                "How will AI impact the Indian financial sector by 2025?",
                "What are the risks and benefits of cryptocurrency adoption?",
                "Analyze the future of remote work and its business implications"
            ]
            for example in examples:
                if st.button(example, key=example):
                    st.rerun()

        else:
            st.error(f"❌ Error: {results.get('message', 'Unknown error')}")
            if results.get('error') == 'rate_limit':
                st.info("💡 This is normal with free Groq tier - wait 60 seconds and try again")

    # Instructions
    with st.expander("ℹ️ How to Use & Deploy", expanded=False):
        st.markdown("""
        ### Usage Instructions
        1. **Enter Groq API Key**: In the sidebar (get one from [Groq Console](https://console.groq.com)).
        2. **Input Query**: Describe what you want to analyze.
        3. **Run Analysis**: Click the button to start the 3-phase multi-agent process.
        4. **View Results**: See structured outputs in columns.
        5. **Download DOCX**: Get a professional Word report if enabled.

        ### Deployment on Streamlit Cloud
        1. Save this as `app.py`.
        2. Create `requirements.txt`:
           ```
           pysqlite3-binary==0.5.4
           chromadb==0.5.11
           streamlit
           crewai==0.186.1
           crewai-tools==0.71.0
           python-docx
           python-dotenv
           rich==13.7.1
           ```
        3. Create `runtime.txt`:
           ```
           python-3.11
           ```
        4. Create `packages.txt`:
           ```
           libsqlite3-dev
           sqlite3
           ```
        5. Push to GitHub repo.
        6. Connect to [Streamlit Cloud](https://share.streamlit.io) and deploy.
        7. For secrets: Add `GROQ_API_KEY` to Streamlit secrets.toml in production.

        ### Notes
        - Analysis takes ~3-4 minutes due to rate limiting.
        - Verbose output is disabled to keep the UI clean.
        - Handles errors gracefully with troubleshooting tips.
        - SQLite patched with pysqlite3-binary to meet ChromaDB requirements.
        """)

if __name__ == "__main__":
    main()
